#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Validate the generated DOCX and report metadata."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = PROJECT_ROOT / "output" / "doc" / "基于MATLAB的机器人端无线控制系统设计报告.docx"
LOGS_DIR = PROJECT_ROOT / "output" / "logs"
RESULTS_DIR = PROJECT_ROOT / "output" / "results"
FIGURES_DIR = PROJECT_ROOT / "output" / "figures"


def count_cjk(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def main():
    if not REPORT_PATH.exists():
        raise FileNotFoundError(REPORT_PATH)
    doc = Document(REPORT_PATH)
    text = "\n".join(p.text for p in doc.paragraphs)
    if "\ufffd" in text:
        raise AssertionError("DOCX contains Unicode replacement character U+FFFD")
    required_figures = [
        "system_architecture.png",
        "motor_control_circuit.png",
        "video_capture_circuit.png",
        "rf_integrated_circuit.png",
        "software_flow.png",
        "simulink_model_top.png",
        "simulink_model_control_subsystem.png",
        "simulink_model_video_rf_subsystems.png",
        "motor_response.png",
        "video_buffer.png",
        "rf_link_ber.png",
        "system_performance_summary.png",
    ]
    missing_figures = [name for name in required_figures if not (FIGURES_DIR / name).exists()]
    if missing_figures:
        raise AssertionError(f"Missing figures: {missing_figures}")
    required_results = [
        "simulation_summary.json",
        "motor_response.csv",
        "video_buffer.csv",
        "rf_link_budget.csv",
        "rf_ber_curve.csv",
        "crossref_index.json",
    ]
    missing_results = [name for name in required_results if not (RESULTS_DIR / name).exists()]
    if missing_results:
        raise AssertionError(f"Missing results: {missing_results}")
    model_path = PROJECT_ROOT / "models" / "robot_wireless_control_system.slx"
    if not model_path.exists():
        raise AssertionError(f"Missing hierarchical Simulink model: {model_path}")
    citations = sorted({int(n) for n in re.findall(r"\[(\d+)\]", text)})
    if len(citations) < 12:
        raise AssertionError(f"Expected at least 12 citations, found {len(citations)}")
    if citations != list(range(1, max(citations) + 1)):
        raise AssertionError(f"Citation numbering is not continuous: {citations}")

    validation_meta = json.loads((LOGS_DIR / "report_generation_validation.json").read_text(encoding="utf-8"))
    body_cjk = validation_meta["body_cjk_chars"]
    if not (9500 <= body_cjk <= 10500):
        raise AssertionError(f"Body CJK character count out of target range: {body_cjk}")

    result = {
        "report": str(REPORT_PATH).replace("\\", "/"),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "body_cjk_chars": body_cjk,
        "citations": len(citations),
        "has_replacement_char": False,
    }
    (LOGS_DIR / "docx_validation.json").write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
