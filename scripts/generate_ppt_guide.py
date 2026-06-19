#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate a DOCX version of the PPT authoring guide from Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = PROJECT_ROOT / "output" / "doc"
MD_PATH = DOC_DIR / "PPT制作说明书.md"
DOCX_PATH = DOC_DIR / "PPT制作说明书.docx"


def set_run_font(run, size: float = 11, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_text(paragraph, text: str, size: float = 11, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_para(doc: Document, text: str, *, indent: bool = False, bullet: bool = False) -> None:
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    add_text(p, text)


def build() -> None:
    if not MD_PATH.exists():
        raise FileNotFoundError(MD_PATH)

    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    for raw in MD_PATH.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, line[2:], size=18, bold=True)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            add_text(p, line[3:], size=14, bold=True)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            add_text(p, line[4:], size=12.5, bold=True)
        elif line.startswith("- "):
            add_para(doc, line[2:], bullet=True)
        elif re.match(r"^\d+\.\s+", line):
            add_para(doc, line)
        elif line.startswith("```"):
            continue
        else:
            add_para(doc, line, indent=True)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
